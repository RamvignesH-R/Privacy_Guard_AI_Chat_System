"""
utils/ocr_extractor.py
Extract text from uploaded documents:
  • PDF  → PyMuPDF (text) or tesseract OCR (scanned)
  • Image (PNG/JPG/TIFF/BMP) → tesseract OCR
  • TXT  → read directly
  • DOCX → python-docx

Install:
    pip install pytesseract Pillow PyMuPDF pdf2image python-docx --break-system-packages
    sudo apt-get install tesseract-ocr poppler-utils   # Linux
    brew install tesseract poppler                      # macOS
"""

import os
import shutil


def _setup_tesseract():
    """Find tesseract binary and configure pytesseract."""
    import pytesseract

    if shutil.which("tesseract"):
        pytesseract.pytesseract.tesseract_cmd = shutil.which("tesseract")
        return pytesseract

    candidates = [
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract",
        "/opt/homebrew/bin/tesseract",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for c in candidates:
        if os.path.isfile(c):
            pytesseract.pytesseract.tesseract_cmd = c
            return pytesseract

    raise EnvironmentError(
        "\n\nTesseract OCR engine not found.\n"
        "Fix with ONE of these:\n"
        "  Ubuntu/Debian : sudo apt-get install tesseract-ocr\n"
        "  macOS         : brew install tesseract\n"
        "  Windows       : https://github.com/UB-Mannheim/tesseract/wiki\n"
        "Then re-run your command."
    )


def extract_text(file_path: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        return _from_txt(file_path)
    elif ext == ".pdf":
        return _from_pdf(file_path)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"):
        return _from_image(file_path)
    elif ext in (".docx", ".doc"):
        return _from_docx(file_path)
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'\n"
            f"Supported: .txt  .pdf  .png  .jpg  .jpeg  .tiff  .bmp  .docx"
        )


def _from_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _from_image(path: str) -> str:
    try:
        from PIL import Image
    except ImportError:
        raise ImportError("Run: pip install Pillow --break-system-packages")

    try:
        pytesseract = _setup_tesseract()
    except ImportError:
        raise ImportError("Run: pip install pytesseract --break-system-packages")

    img = Image.open(path).convert("RGB")

    # Upscale small images for better OCR accuracy
    w, h = img.size
    if w < 1200:
        scale = 1200 / w
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    text = pytesseract.image_to_string(img, config="--psm 6 --oem 3")
    if not text.strip():
        text = pytesseract.image_to_string(img, config="--psm 4 --oem 3")
    return text


def _from_pdf(path: str) -> str:
    try:
        import fitz
        doc  = fitz.open(path)
        full = "\n".join(page.get_text() for page in doc).strip()
        if len(full) > 100:
            return full
    except ImportError:
        pass
    return _pdf_ocr(path)


def _pdf_ocr(path: str) -> str:
    try:
        from pdf2image import convert_from_path
    except ImportError:
        raise ImportError(
            "Run: pip install pdf2image --break-system-packages\n"
            "And: sudo apt-get install poppler-utils"
        )
    pytesseract = _setup_tesseract()
    pages = convert_from_path(path, dpi=200)
    return "\n\n--- Page Break ---\n\n".join(
        pytesseract.image_to_string(p, config="--psm 6 --oem 3") for p in pages
    )


def _from_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Run: pip install python-docx --break-system-packages")

    doc   = Document(path)
    lines = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.append("  |  ".join(cell.text for cell in row.cells))
    return "\n".join(lines)